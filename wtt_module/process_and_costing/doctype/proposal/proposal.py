# Copyright (c) 2023, wtt_module and contributors
# For license information, please see license.txt



import frappe
from frappe.model.document import Document
from frappe import _
from docx import Document as dd
from datetime import date
from frappe.utils.background_jobs import enqueue
import subprocess
import pdfplumber
import traceback
import pyotp
import sys
import smtplib
import traceback

class Proposal(Document):
	def validate(self):
		secret_key=pyotp.random_base32()
		totp = pyotp.TOTP(secret_key, digits=6)
		otp_code = totp.now()
		self.secret_key = otp_code
		
		self.generate_proposal()
		self.send_mail_smtp(otp_code)
		# self.send_mail(otp_code)
	def generate_proposal(self):
		try:
			att=self.attach
			file = frappe.get_doc("File", {'file_url':str(att)})
			file_path = file.get_full_path()
			doc = dd(file_path)
			ar=[]
			pc=0
			rc=0
			for p in doc.paragraphs:
				if 'date_info' in p.text:
					for i in p.runs:
						if(i.text=='date_info'):
							i.text=date.today().strftime("%d-%m-%Y")
				elif 'company_name' in p.text:
					for i in p.runs:
						if(i.text=='company_name'):
							i.text=self.organization
				elif 'CITY' in p.text:
					for i in p.runs:
						if(i.text=='CITY'):
							i.text=self.city
						elif(i.text=='CITY.'):
							i.text=self.city
				elif 'client_name' in p.text:
					for i in p.runs:
						if(i.text=='client_name'):
							i.text=self.client_name
				elif 'proposal_cnt' in p.text:
					for i in p.runs:
						if(i.text=='proposal_cnt'):
							query1=frappe.db.sql("SELECT proposal_no from `tabProposal` where email='"+str(self.email)+"' and capacity='"+str(self.capacity)+"' ",as_dict=1)
							query2=frappe.db.sql("SELECT count(name)as cnt from `tabProposal` where email!='"+str(self.email)+"' ",as_dict=1)
							if(query1):
								pc="{:04d}".format(int(query1[0].proposal_no))
							elif(query2):
								pc=str("{:04d}".format(int(query2[0].cnt)+1))
							else:
								pc="{:04d}".format(1)
							i.text=pc

						if(i.text=='revision_cnt'):
							i.text=str(self.revision)
			for table in doc.tables:
				for row in table.rows:
					for cell in row.cells:
						for paragraph in cell.paragraphs:
							if(paragraph.text=="price_in_docx"):
								amount = frappe.db.get_value("Proposal Document",str(self.capacity),"price_inr")
								if(self.currency=="USD"):
									amount = frappe.db.get_value("Proposal Document",str(self.capacity),"price_usd")
								elif(self.currency=="EUR"):
									amount = frappe.db.get_value("Proposal Document",str(self.capacity),"price_eur")
								paragraph.text=str(amount)
							elif(paragraph.text=="PRICE IN LAKHS"):
								if(self.currency=="USD"):
									paragraph.text = "PRICE IN USD"
								elif(self.currency=="EUR"):
									paragraph.text = "PRICE IN EURO"
			temp_pdf_file_path = "./erp.wttindia.com/public/files/temp.pdf"
			pdf_file_path="./erp.wttindia.com/public/files/Proposal1cb11d.pdf"
			doc.save(temp_pdf_file_path)
			subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_file_path, temp_pdf_file_path])

			self.proposal_no=pc
			self.attach2="/files/Proposal1cb11d.pdf"
		except Exception as e:
			error_message = f"An error occurred: {str(e)} (Line {traceback.extract_tb(e.__traceback__)[0].lineno})"
			frappe.throw(str(error_message))

	
	def send_mail_smtp(self,otp_code): 
		s = smtplib.SMTP_SSL('smtp.gmail.com', 465)
		s.login("noreply@wttint.com", "Aw@rd%25@noreply")
		message = 'Dear '+str(self.client_name)+',\nThis is your One Time Password to access the Proposal Document '+otp_code+'.'
		s.sendmail("noreply@wttint.com", self.email, message)
		s.quit()

	def send_mail(self,otp_code):
		message='<p>Dear '+str(self.client_name)+',<br>This is your One Time Password to access the Proposal Document <b>'+otp_code+'</b>.'
		email_args = {
			"recipients": [self.email],
			"message": message,
			"subject": 'OTP from WTT INTERNATIONAL'
		}

		frappe.sendmail(**email_args)

@frappe.whitelist()
def send_mail():
	# creates SMTP session
	s = smtplib.SMTP_SSL('smtp.gmail.com', 465)
	# start TLS for security
	# s.starttls()	 
	# Authentication
	s.login("noreply@wttint.com", "Aw@rd%25@noreply")		 
	# message to be sent
	message = "test"		 
	# sending the mail
	s.sendmail("noreply@wttint.com", "erp@wttindia.com", message)		 
	# terminating the session
	s.quit()
	return 'test'
@frappe.whitelist()
def get_proposal(att,company,client,capacity,city,currency):
	file = frappe.get_doc("File", {'file_url':str(att)})
	file_path = file.get_full_path()
	doc = dd(file_path)
	ar=[]
	pc=0
	rc=0
	for p in doc.paragraphs:
		if 'date_info' in p.text:
			for i in p.runs:
				if(i.text=='date_info'):
					i.text=date.today().strftime("%d-%m-%Y")
					frappe.msgprint(str(i.text))
		elif 'company_name' in p.text:
			for i in p.runs:
				if(i.text=='company_name'):
					i.text=company
					frappe.msgprint(str(i.text))
		elif 'CITY' in p.text:
			for i in p.runs:
				if(i.text=='CITY'):
					i.text=city
					frappe.msgprint(str(i.text)+".")
				elif(i.text=='CITY.'):
					i.text=city
					frappe.msgprint(str(i.text)+".")
		elif 'client_name' in p.text:
			for i in p.runs:
				if(i.text=='client_name'):
					i.text=client
					frappe.msgprint(str(i.text))
		elif 'proposal_cnt' in p.text:
			for i in p.runs:
				if(i.text=='proposal_cnt'):
					query1=frappe.db.sql("SELECT proposal_no from `tabProposal` where email='erp@wttindia.com' and capacity='"+str(capacity)+"' ",as_dict=1)
					query2=frappe.db.sql("SELECT count(name)as cnt from `tabProposal` where email!='erp@wttindia.com' ",as_dict=1)
					if(query1):
						pc="{:04d}".format(int(query1[0].proposal_no))
					elif(query2):
						pc=str("{:04d}".format(int(query2[0].cnt)+1))
					else:
						pc="{:04d}".format(1)
					i.text=pc
					frappe.msgprint(str(i.text))

				if(i.text=='revision_cnt'):
					i.text='00'
					frappe.msgprint(str(i.text))

			# 			if "price_in_docx" in paragraph.text:
			# 				frappe.msgprint(paragraph.text)
							# paragraph.text = paragraph.text.replace(old_text, new_text)
		ar.append(p.text)
	for table in doc.tables:
		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					if(paragraph.text=="price_in_docx"):
						amount = frappe.db.get_value("Proposal Document",str(capacity),"price_inr")
						if(currency=="USD"):
							amount = frappe.db.get_value("Proposal Document",str(capacity),"price_usd")
						elif(currency=="EUR"):
							amount = frappe.db.get_value("Proposal Document",str(capacity),"price_eur")
						paragraph.text=str(amount)

						frappe.msgprint(str(paragraph.text))
	return "test"



@frappe.whitelist(allow_guest=True)
def get_attachment(capacity):
	ar=frappe.db.get_value("Proposal Document",capacity,"attachment")
	return ar